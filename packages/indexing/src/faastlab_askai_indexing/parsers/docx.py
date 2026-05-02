"""DOCX parser — python-docx for paragraph + heading extraction."""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument

from faastlab_askai_indexing.parsers.base import (
    ParsedBlock,
    ParsedDocument,
    ParserError,
)

_DOCX_MIME_TYPES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)


class DocxParser:
    """DOCX → ParsedDocument."""

    @property
    def supported_content_types(self) -> tuple[str, ...]:
        return _DOCX_MIME_TYPES

    def parse(self, data: bytes, *, filename: str | None = None) -> ParsedDocument:
        try:
            doc = DocxDocument(BytesIO(data))
        except Exception as exc:  # noqa: BLE001
            raise ParserError(f"python-docx failed to open DOCX: {exc}") from exc

        blocks: list[ParsedBlock] = []
        char_offset = 0
        section_path: list[str] = []
        title: str | None = None

        for paragraph in doc.paragraphs:
            text_value = (paragraph.text or "").strip()
            if not text_value:
                continue

            style_name = (paragraph.style.name or "").strip() if paragraph.style else ""
            block_type: str = "paragraph"
            if style_name.startswith("Heading"):
                block_type = "heading"
                if not title:
                    title = text_value
                # Track depth via heading level for section_path.
                level = _heading_level(style_name)
                if level is not None:
                    section_path = section_path[: max(level - 1, 0)]
                    section_path.append(text_value)
            elif style_name == "Title":
                block_type = "title"
                title = title or text_value
            elif style_name.startswith("List"):
                block_type = "list_item"

            start = char_offset
            end = char_offset + len(text_value)
            blocks.append(
                ParsedBlock(
                    text=text_value,
                    block_type=block_type,  # type: ignore[arg-type]
                    section_path=" / ".join(section_path) if section_path else None,
                    char_start=start,
                    char_end=end,
                )
            )
            char_offset = end + 2

        if (not title) and filename:
            title = filename.rsplit("/", 1)[-1].rsplit(".", 1)[0]

        return ParsedDocument(
            title=title,
            blocks=blocks,
            metadata={"parser": "python-docx"},
        )


def _heading_level(style_name: str) -> int | None:
    """`Heading 1` → 1, `Heading 2` → 2, …"""
    parts = style_name.split()
    if len(parts) == 2 and parts[1].isdigit():
        return int(parts[1])
    return None
